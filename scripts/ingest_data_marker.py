# scripts/ingest_data_marker.py (Versión para DOCX con Marker)
import os
import pathlib
import numpy as np
import requests
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from transformers import AutoTokenizer
import psycopg2
from pgvector.psycopg2 import register_vector

print("Iniciando script de ingesta para archivos DOCX usando Marker...")

# --- 1. Carga de Configuración ---
load_dotenv(override=True)
EMBEDDING_API_ENDPOINT = os.environ.get("EMBEDDING_API_ENDPOINT")
DB_CONNECTION_STRING = os.environ.get("DB_CONNECTION_STRING")
TOKENIZER_NAME = "sentence-transformers/all-MiniLM-L6-v2"

if not EMBEDDING_API_ENDPOINT or not DB_CONNECTION_STRING:
    print("Error: Falta EMBEDDING_API_ENDPOINT o DB_CONNECTION_STRING en el .env.")
    exit()

# --- 2. Importación de Marker con manejo de errores ---
try:
    from marker import Marker
    print("Marker importado correctamente.")
    marker_available = True
except ImportError as e:
    print(f"Error importando Marker: {e}")
    print("Instalando dependencias de Marker...")
    try:
        import subprocess
        subprocess.check_call(["pip", "install", "--upgrade", "marker-pdf"])
        from marker import Marker
        print("Marker instalado e importado correctamente.")
        marker_available = True
    except Exception as install_error:
        print(f"No se pudo instalar Marker: {install_error}")
        print("Usando python-docx como alternativa...")
        marker_available = False

# --- 3. Función para Obtener Embeddings con Requests ---
def get_embedding_from_vllm(text: str, endpoint: str) -> np.ndarray:
    """Obtiene el embedding de un texto haciendo una llamada HTTP directa al endpoint vLLM."""
    try:
        response = requests.post(endpoint, json={"text": text}, timeout=60)
        response.raise_for_status()
        data = response.json()
        return np.array(data["embedding"])
    except Exception as e:
        print(f"Error llamando al endpoint de embedding vLLM ({endpoint}): {e}")
        return None

# --- 4. Conexión a PostgreSQL ---
def db_connection():
    """Establece y devuelve una conexión a la base de datos PostgreSQL."""
    try:
        conn = psycopg2.connect(DB_CONNECTION_STRING)
        register_vector(conn)
        print("Conexión a PostgreSQL establecida.")
        return conn
    except Exception as e:
        print(f"Error conectando a PostgreSQL: {e}")
        exit()

# --- 5. Definiciones y Carga del Tokenizador ---
SCRIPT_DIR = pathlib.Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DATA_DIR = PROJECT_ROOT / "data" / "doc"

try:
    filenames = [f for f in os.listdir(DATA_DIR) if f.lower().endswith(".docx")]
    if not filenames:
        print(f"No se encontraron archivos DOCX en: {DATA_DIR}"); exit()
    print(f"Archivos DOCX a procesar: {filenames}")
except Exception as e:
    print(f"Error listando archivos DOCX: {e}"); exit()

try:
    tokenizer = AutoTokenizer.from_pretrained(TOKENIZER_NAME)
    print(f"Tokenizador '{TOKENIZER_NAME}' cargado.")
except Exception as e:
    print(f"Error al cargar tokenizador '{TOKENIZER_NAME}': {e}"); exit()

# --- 6. Creación de Tabla (si no existe) ---
def create_documents_table_if_not_exists(conn):
    """Crea la tabla 'documents_wifi' si no existe."""
    with conn.cursor() as cur:
        cur.execute("""
        CREATE TABLE IF NOT EXISTS documents_security (
            id TEXT PRIMARY KEY,
            text_content TEXT,          
            embedding VECTOR(384)
        );
        """)
        conn.commit()
        print("Tabla 'documents_wifi' verificada/creada.")

# --- 7. Función para Extraer Texto con Marker ---
def extract_text_with_marker(file_path):
    """Extrae texto usando Marker para mejor calidad."""
    try:
        if not marker_available:
            return None
            
        # Inicializar Marker
        marker = Marker()
        
        # Procesar el documento
        doc = marker.convert(str(file_path))
        
        # Extraer texto
        text_content = ""
        for page in doc.pages:
            for block in page.blocks:
                if hasattr(block, 'text'):
                    text_content += block.text + "\n"
                elif hasattr(block, 'table'):
                    # Procesar tablas
                    for row in block.table.rows:
                        row_text = []
                        for cell in row.cells:
                            if hasattr(cell, 'text'):
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_content += " | ".join(row_text) + "\n"
        
        return text_content.strip()
        
    except Exception as e:
        print(f"Error extrayendo texto con Marker: {e}")
        return None

# --- 8. Función de respaldo con python-docx ---
def extract_text_with_docx(file_path):
    """Extrae texto usando python-docx como respaldo."""
    try:
        from docx import Document
        doc = Document(file_path)
        all_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_text.append(paragraph.text)
        
        # Procesar tablas
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
            
            if table_text:
                table_content = "\n".join(table_text)
                all_text.append(f"Tabla:\n{table_content}")
        
        return "\n\n".join(all_text)
    except Exception as e:
        print(f"Error extrayendo texto con python-docx: {e}")
        return None

# --- 9. Procesamiento e Ingesta ---
def main():
    """Función principal que orquesta la ingesta de archivos DOCX."""
    conn_pg = db_connection()
    create_documents_table_if_not_exists(conn_pg)

    text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
        tokenizer=tokenizer, chunk_size=768, chunk_overlap=75
    )

    for filename in filenames:
        full_file_path = DATA_DIR / filename
        print(f"\nProcesando archivo: {full_file_path}...")
        
        # Intentar extraer texto con Marker primero
        document_text = None
        if marker_available:
            print("  Intentando extracción con Marker...")
            document_text = extract_text_with_marker(full_file_path)
        
        # Si Marker falla, usar python-docx como respaldo
        if not document_text:
            print("  Usando python-docx como respaldo...")
            document_text = extract_text_with_docx(full_file_path)
        
        if not document_text:
            print(f"  No se pudo extraer texto del archivo '{filename}'. Saltando...")
            continue
        
        # Dividir el texto en chunks
        text_chunks = text_splitter.split_text(document_text)
        print(f"  Total de {len(text_chunks)} fragmentos generados para '{filename}'.")

        # Generar embeddings e insertar en la base de datos
        file_chunks_for_db = []
        for i, chunk_text in enumerate(text_chunks):
            if not chunk_text or not chunk_text.strip(): 
                continue

            chunk_id = f"{filename}-{(i + 1)}"
            embedding_vector = get_embedding_from_vllm(chunk_text, EMBEDDING_API_ENDPOINT)
            
            if embedding_vector is not None:
                file_chunks_for_db.append({"id": chunk_id, "text_content": chunk_text, "embedding": embedding_vector})
            else:
                print(f"    Advertencia: Se saltará el chunk '{chunk_id}' debido a un error de embedding.")
        
        # Inserta los chunks en la base de datos
        if file_chunks_for_db:
            try:
                with conn_pg.cursor() as cur:
                    insert_query = "INSERT INTO documents_security (id, text_content, embedding) VALUES (%s, %s, %s) ON CONFLICT (id) DO UPDATE SET text_content = EXCLUDED.text_content, embedding = EXCLUDED.embedding;"
                    data_to_insert = [(c["id"], c["text_content"], c["embedding"]) for c in file_chunks_for_db]
                    cur.executemany(insert_query, data_to_insert)
                    conn_pg.commit()
                print(f"  Archivo '{filename}' procesado: {len(file_chunks_for_db)} fragmentos insertados/actualizados.")
            except Exception as e:
                print(f"  Error al insertar datos para '{filename}': {e}")
                if conn_pg: conn_pg.rollback()
        else:
            print(f"  No se generaron fragmentos válidos para '{filename}'.")
            
    if conn_pg:
        conn_pg.close()
        print("\nConexión a PostgreSQL cerrada.")
    print("\nProceso de ingestión a PostgreSQL completado.")

if __name__ == "__main__":
    main() 