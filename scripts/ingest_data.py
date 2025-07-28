# scripts/ingest_data.py (Versión para DOCX con python-docx)
import os
import pathlib
import numpy as np
import requests
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from transformers import AutoTokenizer
import psycopg2
from pgvector.psycopg2 import register_vector
from docx import Document  # Importar la librería python-docx para procesar DOCX
from marker.converters.pdf import PdfConverter
from marker.models import create_model_dict
from marker.output import text_from_rendered

print("Iniciando script de ingesta para archivos DOCX...")

# --- 1. Carga de Configuración ---
load_dotenv(override=True)
EMBEDDING_API_ENDPOINT = os.environ.get("EMBEDDING_API_ENDPOINT")
DB_CONNECTION_STRING = os.environ.get("DB_CONNECTION_STRING")
TOKENIZER_NAME = "sentence-transformers/all-MiniLM-L6-v2"

SCRIPT_DIR = pathlib.Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DATA_DIR = PROJECT_ROOT / "data" / "pdfs"

if not EMBEDDING_API_ENDPOINT or not DB_CONNECTION_STRING:
    print("Error: Falta EMBEDDING_API_ENDPOINT o DB_CONNECTION_STRING en el .env.")
    exit()

# --- 2. Función para Obtener Embeddings con Requests ---
def get_embedding_from_vllm(text: str, endpoint: str) -> np.ndarray:
    """Obtiene el embedding de un texto haciendo una llamada HTTP directa al endpoint vLLM."""
    try:
        response = requests.post(endpoint, json={"text": text}, timeout=60)
        response.raise_for_status()
        data = response.json()
        return np.array(data["embedding"])
    except Exception as e:
        print(f"Error llamando al endpoint de embedding vLLM ({endpoint}): {e}")
        return None

# --- 3. Conexión a PostgreSQL ---
def db_connection():
    """Establece y devuelve una conexión a la base de datos PostgreSQL."""
    try:
        conn = psycopg2.connect(DB_CONNECTION_STRING)
        register_vector(conn)
        print("Conexión a PostgreSQL establecida.")
        return conn
    except Exception as e:
        print(f"Error conectando a PostgreSQL: {e}")
        exit()


try:
    filenames = [f for f in os.listdir(DATA_DIR) if f.lower().endswith(".docx")]
    if not filenames:
        print(f"No se encontraron archivos DOCX en: {DATA_DIR}"); exit()
    print(f"Archivos DOCX a procesar: {filenames}")
except Exception as e:
    print(f"Error listando archivos DOCX: {e}"); exit()

try:
    tokenizer = AutoTokenizer.from_pretrained(TOKENIZER_NAME)
    print(f"Tokenizador '{TOKENIZER_NAME}' cargado.")
except Exception as e:
    print(f"Error al cargar tokenizador '{TOKENIZER_NAME}': {e}"); exit()

# --- 5. Creación de Tabla (si no existe) ---
def create_documents_table_if_not_exists(conn):
    """Crea la tabla 'documents_wifi' si no existe."""
    with conn.cursor() as cur:
        cur.execute("""
        CREATE TABLE IF NOT EXISTS documents_wifi (
            id TEXT PRIMARY KEY,
            text_content TEXT,          
            embedding VECTOR(384)
        );
        """)
        conn.commit()
        print("Tabla 'documents_wifi' verificada/creada.")

# --- 6. Función para Extraer Texto de DOCX ---
def extract_text_from_docx(file_path):
    """Extrae todo el texto de un archivo DOCX, incluyendo tablas."""
    try:
        doc = Document(file_path)
        all_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_text.append(paragraph.text)
        
        # Procesar tablas
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
            
            if table_text:
                table_content = "\n".join(table_text)
                all_text.append(f"Tabla:\n{table_content}")
        
        return "\n\n".join(all_text)
    except Exception as e:
        print(f"Error extrayendo texto del archivo DOCX: {e}")
        return None

# --- 7. Procesamiento e Ingesta ---
def main():
    """Función principal que orquesta la ingesta de archivos DOCX."""
    conn_pg = db_connection()
    create_documents_table_if_not_exists(conn_pg)

    text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
        tokenizer=tokenizer, chunk_size=512, chunk_overlap=50
    )

    for filename in filenames:
        full_file_path = DATA_DIR / filename
        print(f"\nProcesando archivo: {full_file_path}...")
        
        # Extrae el texto del PDF.
        try:
            rendered = converter(str(full_file_path))
            md_text, _, images = text_from_rendered(rendered)
            if not md_text.strip(): # Si el PDF no tiene texto extraíble.
                print(f"  Advertencia: No se extrajo texto de '{filename}' usando marker.")
                continue
        except Exception as e:
            print(f"  Error al leer/convertir PDF '{filename}' con marker: {e}")
            continue

        # Divide el texto extraído en fragmentos (chunks).
        text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
            tokenizer=tokenizer, chunk_size=512, chunk_overlap=50
        )
        texts_docs = text_splitter.create_documents([md_text])
        
        file_chunks_for_db = [] # Lista para almacenar los chunks listos para la BD.
        print(f"  Dividido en {len(texts_docs)} fragmentos.")

        # Procesa cada fragmento: genera ID, obtiene texto y genera embedding.
        for i, text_doc_obj in enumerate(texts_docs):
            chunk_id = f"{filename}-{(i + 1)}" # ID único para el chunk.
            chunk_text = text_doc_obj.page_content
            embedding_vector = get_embedding_from_vllm(chunk_text, EMBEDDING_API_ENDPOINT)
            
            if embedding_vector is not None:
                file_chunks_for_db.append({"id": chunk_id, "text_content": chunk_text, "embedding": embedding_vector})
            else:
                print(f"    Advertencia: Se saltará el chunk '{chunk_id}' debido a un error de embedding.")
        
        # Inserta los chunks en la base de datos
        if file_chunks_for_db:
            try:
                with conn_pg.cursor() as cur:
                    insert_query = "INSERT INTO documents_wifi (id, text_content, embedding) VALUES (%s, %s, %s) ON CONFLICT (id) DO UPDATE SET text_content = EXCLUDED.text_content, embedding = EXCLUDED.embedding;"
                    data_to_insert = [(c["id"], c["text_content"], c["embedding"]) for c in file_chunks_for_db]
                    cur.executemany(insert_query, data_to_insert)
                    conn_pg.commit()
                print(f"  Archivo '{filename}' procesado: {len(file_chunks_for_db)} fragmentos insertados/actualizados.")
            except Exception as e:
                print(f"  Error al insertar datos para '{filename}': {e}")
                if conn_pg: conn_pg.rollback()
        else:
            print(f"  No se generaron fragmentos válidos para '{filename}'.")
            
    if conn_pg:
        conn_pg.close()
        print("\nConexión a PostgreSQL cerrada.")
    print("\nProceso de ingestión a PostgreSQL completado.")

if __name__ == "__main__":
    main()